from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from connectionExecution import connectionResults
from fillFooter import fillingFooterFunction
from merge import combineParts

# Assuming you have a function connectionResults(80) to retrieve data
headerFooterResult, rowResult = connectionResults(10)
headerFooterList = headerFooterResult[0]

finalDataList = []

header_data = [
    "", "", "345345", "XCXC",
    str(headerFooterList[5]), str(headerFooterList[5]),
    "", str(headerFooterList[1]), "G023", str(headerFooterList[2])
]

for rRow in rowResult:
    internalList = [str(rRow[7]), str(rRow[5]), str(rRow[9]), str(rRow[5]), str(
        rRow[6]) + str(rRow[7]), str(rRow[1]), str(rRow[3]), str(rRow[2]), "1"]
    finalDataList.append(internalList)

chunk_size = 7
data_chunks = [finalDataList[i:i + chunk_size]
               for i in range(0, len(finalDataList), chunk_size)]

file_names = []
for i, chunk in enumerate(data_chunks):
    document = Document('pt1.docx')
    style = document.styles['Normal']
    font = style.font
    font.name = 'Cascadia Code'
    font.complex_script = True
    font.rtl = True
    font.size = Pt(8)
    if len(document.tables) > 1:
        table = document.tables[1]
        num_rows = len(table.rows)
        if num_rows >= len(chunk):
            for j in range(len(chunk)):
                for k in range(len(chunk[j])):
                    cell = table.cell(j, k)
                    cell.text = chunk[j][k]
        else:
            print(
                f"Table 1 in 'pt1.docx' does not have enough rows for chunk {i + 1}.")

    if len(document.tables) > 0:
        header_table = document.tables[0]
        if len(header_table.rows) * len(header_table.columns) == len(header_data):
            for j, data in enumerate(header_data):
                row_index = j // len(header_table.columns)
                col_index = j % len(header_table.columns)
                header_cell = header_table.cell(row_index, col_index)
                header_cell.text = data
                for paragraph in header_cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            print(
                "The number of cells in the table does not match the length of header_data.")

    file_name = f'main{i + 1}.docx'
    document.save(file_name)
    file_names.append(file_name)

fillingFooterFunction()
file_names.append('footer.docx')
combineParts(file_names)
print("Generated files:", file_names)
