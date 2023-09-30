from docx import Document
from docx.shared import Pt
from connectionExecution import connectionResults



def fillingFooterFunction():
    headerFooterResult, rowResult = connectionResults(80)
    headerFooterList = headerFooterResult[0]

    document = Document('pt2.docx')

    style = document.styles['Normal']
    font = style.font
    font.name = 'Cascadia Code'
    font.complex_script = True
    font.rtl = True
    font.size = Pt(8)


    footer_data = [str(headerFooterList[8]), str(headerFooterList[9]), str(
        headerFooterList[10]), str(headerFooterList[11]), str(headerFooterList[12])]


    if len(document.tables) > 0:
        table = document.tables[0]
        for i in range(len(footer_data)):
            for j in range(len(table.columns)):
                cell = table.cell(i, j)
                cell.text = footer_data[i]

    document.save('footer.docx')
