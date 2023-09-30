from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from connectionExecution import connectionResults
from fillFooter import fillingFooterFunction
headerFooterResult, rowResult = connectionResults(80)
headerFooterList = headerFooterResult[0]
finalDataList = []

counter = 1
for rRow in rowResult:
    internalList = [str(rRow[7]), str(rRow[5]), str(rRow[9]), str(rRow[5]), str(
        rRow[6]) + str(rRow[7]), str(rRow[1]), str(rRow[3]), str(rRow[2]), "1"]

    counter = counter + 1
    if counter == 9:
        break
    finalDataList.append(internalList)

document = Document('pt1.docx')

style = document.styles['Normal']
font = style.font
font.name = 'Cascadia Code'
font.complex_script = True
font.rtl = True
font.size = Pt(8)

dummy_data = finalDataList

header_data = [
    "", "", "345345", "XCXC",
    str(headerFooterList[5]), str(headerFooterList[5]),
    "", str(headerFooterList[2]), str(headerFooterList[1])
]

if len(document.tables) > 0:
    table = document.tables[0]
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Iterate over all 9 cells in the table
    for i in range(len(header_data)):
        row = i // len(table.columns)
        col = i % len(table.columns)
        cell = table.cell(row, col)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.text = header_data[i]

    # Set the text for the merged cell in the last row
    merged_cell = table.cell(4, 0)  # Assuming it's the first cell in the last row
    merged_cell.text = header_data[7]
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


if len(document.tables) > 1:
    table = document.tables[1]
    num_rows = len(table.rows)
    num_cols = len(table.columns)
    if num_rows >= len(dummy_data) and num_cols >= len(dummy_data[0]):
        for i in range(len(dummy_data)):
            for j in range(len(dummy_data[i])):
                cell = table.cell(i, j)
                cell.text = dummy_data[i][j]
    else:
        print("Table 1 does not have enough rows or columns for the dummy data.")

document.save('main1.docx')
fillingFooterFunction()
